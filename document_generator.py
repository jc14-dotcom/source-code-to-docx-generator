import os
import logging
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from models import FileInfo, ScanResult, DocumentStats
from secrets_detector import SecretDetector
from framework_profiles import FrameworkProfile

logger = logging.getLogger("docgen")

LINES_PER_PAGE = 56
FRONT_MATTER_PAGES = 3
METADATA_LINES = 9
SEPARATOR_LINES = 2
SECTION_HEADER_LINES = 2
SOURCE_CHARS_PER_LINE = 100
TOC_CHARS_PER_LINE = 78


class _PaginationCursor:
    """Deterministic line-based pagination model used for TOC estimates."""

    def __init__(self, page: int, lines_used: int = 0) -> None:
        self.page = page
        self.lines_used = lines_used

    def consume(self, lines: int) -> None:
        remaining = max(0, lines)
        while remaining:
            if self.lines_used == LINES_PER_PAGE:
                self.page += 1
                self.lines_used = 0

            available = LINES_PER_PAGE - self.lines_used
            consumed = min(remaining, available)
            self.lines_used += consumed
            remaining -= consumed

            if remaining:
                self.page += 1
                self.lines_used = 0

    def explicit_page_break(self) -> None:
        self.page += 1
        self.lines_used = 0


def _set_cell_shading(cell, color: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)


class DocumentGenerator:
    def __init__(self) -> None:
        self.doc: Optional[Document] = None
        self.detector = SecretDetector()
        self.total_secrets_redacted = 0
        self._content_cache: Dict[Path, str] = {}

    def _setup_document(self) -> None:
        self.doc = Document()

        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(14.0)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    def _setup_styles(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Consolas"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        heading_configs = {
            "Heading 1": Pt(16),
            "Heading 2": Pt(14),
            "Heading 3": Pt(12),
        }
        for name, size in heading_configs.items():
            hs = self.doc.styles[name]
            hs.font.name = "Consolas"
            hs.font.size = size
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)

        code_style = self.doc.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = self.doc.styles["Normal"]
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(0)
        code_style.paragraph_format.space_after = Pt(0)
        code_style.paragraph_format.line_spacing = 1.0

    def _add_header_footer(self, system_name: str) -> None:
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(system_name)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.bold = True

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Source Code Documentation")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        fp.add_run("    |    ")
        _add_page_number(fp)

    def _add_horizontal_line(self) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _add_metadata_table(self, info: FileInfo) -> None:
        table = self.doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        rows_data = [
            ("Filename:", info.filename),
            ("Directory:", info.directory),
            ("Language:", info.language),
            ("File Type:", info.file_type),
            ("File Size:", self._format_size(info.size_bytes)),
            ("Last Modified:", info.last_modified.strftime("%Y-%m-%d %H:%M:%S")),
        ]

        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            cell_label = row.cells[0]
            cell_value = row.cells[1]

            cell_label.width = Inches(1.5)
            cell_value.width = Inches(5.0)

            lp = cell_label.paragraphs[0]
            lp.clear()
            lr = lp.add_run(label)
            lr.font.name = "Consolas"
            lr.font.size = Pt(10)
            lr.font.bold = True

            vp = cell_value.paragraphs[0]
            vp.clear()
            vr = vp.add_run(value)
            vr.font.name = "Consolas"
            vr.font.size = Pt(10)

        self.doc.add_paragraph()

    def _format_size(self, size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} bytes"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.2f} MB"

    @staticmethod
    def _sanitize_xml(text: str) -> str:
        return "".join(
            ch for ch in text
            if ch in ("\t", "\n", "\r") or (ord(ch) >= 32 and ord(ch) != 0xFFFE and ord(ch) != 0xFFFF)
        )

    def _add_file_entry(self, info: FileInfo, content: str) -> None:
        self._add_horizontal_line()
        self._add_metadata_table(info)

        # One paragraph with line breaks is considerably cheaper than creating
        # one paragraph for every source line. Word still renders the source
        # line-by-line, while python-docx creates far fewer XML elements.
        p = self.doc.add_paragraph(style="Source Code")
        p.paragraph_format.keep_together = False
        run = p.add_run(self._sanitize_xml(content))
        run.font.name = "Consolas"
        run.font.size = Pt(9)

        self._add_horizontal_line()

    def _read_file_content(self, info: FileInfo) -> str:
        """Read a source file once per generation and reuse its content."""
        if info.path not in self._content_cache:
            self._content_cache[info.path] = info.path.read_text(
                encoding="utf-8",
                errors="replace",
            )
        return self._content_cache[info.path]

    def _validate_docx_secrets(self, path: Path) -> None:
        """Fail closed if known secrets remain in the generated DOCX archive."""
        findings: Dict[str, int] = {}

        with zipfile.ZipFile(path) as archive:
            for member in archive.namelist():
                if not (member.startswith("word/") and member.endswith(".xml")):
                    continue

                try:
                    root = ET.fromstring(archive.read(member))
                except ET.ParseError:
                    continue

                text = "".join(root.itertext())
                for label, count in self.detector.find_labels(text).items():
                    findings[label] = findings.get(label, 0) + count

        if findings:
            categories = ", ".join(sorted(findings))
            raise ValueError(
                "Generated DOCX failed the secret safety check "
                f"({categories}). The output was not published."
            )

    def _add_title_page(self, system_name: str, profile: FrameworkProfile) -> None:
        for _ in range(4):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(system_name)
        run.font.name = "Consolas"
        run.font.size = Pt(24)
        run.font.bold = True

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Complete Source Code Documentation")
        run.font.name = "Consolas"
        run.font.size = Pt(18)
        run.font.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        details = [
            ("Purpose", "Software Copyright Registration"),
            ("Framework", profile.name),
            ("Programming Language", profile.language),
            ("Database", profile.database),
            ("Generated Date", datetime.now().strftime("%B %d, %Y")),
            ("Prepared By", "[Your Name]"),
        ]

        for label, value in details:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}: ")
            run.font.name = "Consolas"
            run.font.size = Pt(12)
            run.font.bold = True
            run = p.add_run(value)
            run.font.name = "Consolas"
            run.font.size = Pt(12)

        self.doc.add_page_break()

    def _add_summary_page(self, stats: DocumentStats, total_folders: int, profile: FrameworkProfile) -> None:
        p = self.doc.add_paragraph("System Summary", style="Heading 1")

        summary_data = [
            ("Total Folders Scanned", str(total_folders)),
            ("Total Files Scanned", str(stats.total_files)),
        ]

        for key, label in profile.stat_labels.items():
            count = stats.get(key)
            if count > 0:
                summary_data.append((label, str(count)))

        table = self.doc.add_table(rows=len(summary_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (label, value) in enumerate(summary_data):
            row = table.rows[i]
            cell_label = row.cells[0]
            cell_value = row.cells[1]

            cell_label.width = Inches(3.0)
            cell_value.width = Inches(2.0)

            if i == 0:
                _set_cell_shading(cell_label, "2B579A")
                _set_cell_shading(cell_value, "2B579A")
                lp = cell_label.paragraphs[0]
                lr = lp.add_run(label)
                lr.font.name = "Consolas"
                lr.font.size = Pt(10)
                lr.font.bold = True
                lr.font.color.rgb = RGBColor(255, 255, 255)
                vp = cell_value.paragraphs[0]
                vr = vp.add_run(value)
                vr.font.name = "Consolas"
                vr.font.size = Pt(10)
                vr.font.bold = True
                vr.font.color.rgb = RGBColor(255, 255, 255)
            else:
                lp = cell_label.paragraphs[0]
                lp.clear()
                lr = lp.add_run(label)
                lr.font.name = "Consolas"
                lr.font.size = Pt(10)

                vp = cell_value.paragraphs[0]
                vp.clear()
                vr = vp.add_run(value)
                vr.font.name = "Consolas"
                vr.font.size = Pt(10)
                vr.font.bold = True

        self.doc.add_page_break()

    @staticmethod
    def _ceil_div(value: int, divisor: int) -> int:
        return max(1, (value + divisor - 1) // divisor)

    @classmethod
    def _estimate_wrapped_lines(cls, content: str) -> int:
        lines = content.split("\n")
        return sum(
            cls._ceil_div(max(1, len(line)), SOURCE_CHARS_PER_LINE)
            for line in lines
        )

    def _count_file_lines(self, info: FileInfo) -> int:
        try:
            content = self._read_file_content(info)
            code_lines = self._estimate_wrapped_lines(content)
            return METADATA_LINES + code_lines + SEPARATOR_LINES
        except Exception:
            return METADATA_LINES + 10 + SEPARATOR_LINES

    def _toc_specs(
        self,
        files_by_part: Dict[str, List[FileInfo]],
        profile: FrameworkProfile,
    ) -> List[Tuple[int, str]]:
        specs: List[Tuple[int, str]] = []

        for part_num, part_title in profile.parts:
            part_files = files_by_part.get(part_num, [])
            if not part_files:
                continue

            specs.append((1, f"PART {part_num} - {part_title}"))
            sections: Dict[str, List[FileInfo]] = {}
            for file_info in part_files:
                sections.setdefault(file_info.section, []).append(file_info)
            for section_name in sections:
                if section_name != part_title:
                    specs.append((2, section_name))

        if files_by_part.get("Appendix"):
            specs.append((1, "Appendix"))

        return specs

    @staticmethod
    def _estimate_toc_lines(specs: List[Tuple[int, str]]) -> int:
        header_lines = 1
        return header_lines + sum(
            max(
                1,
                (len(text) + (level - 1) * 4 + TOC_CHARS_PER_LINE - 1)
                // TOC_CHARS_PER_LINE,
            )
            for level, text in specs
        )

    def _calculate_toc(
        self,
        files_by_part: Dict[str, List[FileInfo]],
        profile: FrameworkProfile,
    ) -> List[Tuple[int, str, int]]:
        toc_entries: List[Tuple[int, str, int]] = []
        specs = self._toc_specs(files_by_part, profile)
        toc_pages = self._ceil_div(self._estimate_toc_lines(specs), LINES_PER_PAGE)
        cursor = _PaginationCursor(page=FRONT_MATTER_PAGES + toc_pages)

        for part_num, part_title in profile.parts:
            part_files = files_by_part.get(part_num, [])
            if not part_files:
                continue

            toc_entries.append((1, f"PART {part_num} - {part_title}", cursor.page))
            cursor.explicit_page_break()

            sections: Dict[str, List[FileInfo]] = {}
            for f in part_files:
                sections.setdefault(f.section, []).append(f)

            for section_name, section_files in sections.items():
                if section_name != part_title:
                    toc_entries.append((2, section_name, cursor.page))
                    cursor.consume(SECTION_HEADER_LINES)

                for file_info in section_files:
                    cursor.consume(self._count_file_lines(file_info))

            cursor.explicit_page_break()

        appendix_files = files_by_part.get("Appendix", [])
        if appendix_files:
            toc_entries.append((1, "Appendix", cursor.page))
            cursor.explicit_page_break()
            for file_info in appendix_files:
                cursor.consume(self._count_file_lines(file_info))

        return toc_entries

    def _add_toc(self, toc_entries: List[Tuple[int, str, int]]) -> None:
        p = self.doc.add_paragraph(
            "Table of Contents (estimated page numbers)",
            style="Heading 1",
        )
        p.paragraph_format.space_after = Pt(18)

        if not toc_entries:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[No sections found]")
            run.font.name = "Consolas"
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            self.doc.add_page_break()
            return

        for level, text, page_num in toc_entries:
            indent = "    " * (level - 1)
            label = f"{indent}{text}"
            dots_needed = max(3, 70 - len(label) - len(str(page_num)))
            dots = "." * dots_needed
            line = f"{label} {dots} {page_num}"

            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(11)
            if level == 1:
                run.font.bold = True

        self.doc.add_page_break()

    def _add_part_header(self, part_number: str, part_title: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(f"PART {part_number}")
        run.font.name = "Consolas"
        run.font.size = Pt(18)
        run.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(part_title)
        run.font.name = "Consolas"
        run.font.size = Pt(14)
        run.font.bold = True

        self.doc.add_page_break()

    def _add_section_header(self, section_title: str) -> None:
        p = self.doc.add_paragraph(section_title, style="Heading 2")
        p.paragraph_format.space_before = Pt(18)

    def _group_files(self, scan_result: ScanResult) -> Dict[str, List[FileInfo]]:
        files_by_part: Dict[str, List[FileInfo]] = {}
        for info in scan_result.files:
            files_by_part.setdefault(info.part, []).append(info)
        return files_by_part

    def _write_parts(
        self,
        files_by_part: Dict[str, List[FileInfo]],
        profile: FrameworkProfile,
        scan_result: ScanResult,
    ) -> None:
        for part_num, part_title in profile.parts:
            part_files = files_by_part.get(part_num, [])
            if not part_files:
                continue

            self._add_part_header(part_num, part_title)

            sections: Dict[str, List[FileInfo]] = {}
            for f in part_files:
                sections.setdefault(f.section, []).append(f)

            for section_name, section_files in sections.items():
                if section_name != part_title:
                    self._add_section_header(section_name)

                for info in section_files:
                    try:
                        content = self._read_file_content(info)
                    except Exception as e:
                        logger.warning(f"Cannot read {info.relative_path}: {e}")
                        scan_result.errors.append(f"Read error: {info.relative_path}: {e}")
                        continue

                    redacted_content, secrets_count = self.detector.scan_and_redact(content)
                    self.total_secrets_redacted += secrets_count

                    if secrets_count > 0:
                        logger.info(f"Redacted {secrets_count} secret(s) in {info.relative_path}")

                    self._add_file_entry(info, redacted_content)

            self.doc.add_page_break()

        appendix_files = files_by_part.get("Appendix", [])
        if appendix_files:
            self._add_part_header("App", "Appendix")
            for info in appendix_files:
                try:
                    content = self._read_file_content(info)
                except Exception as e:
                    logger.warning(f"Cannot read {info.relative_path}: {e}")
                    continue

                redacted_content, secrets_count = self.detector.scan_and_redact(content)
                self.total_secrets_redacted += secrets_count
                self._add_file_entry(info, redacted_content)

            self.doc.add_page_break()

    def generate(
        self,
        scan_result: ScanResult,
        output_path: Path,
        profile: FrameworkProfile,
        system_name: str,
    ) -> None:
        logger.info("Generating document...")

        self.total_secrets_redacted = 0
        self._content_cache.clear()
        self._setup_document()
        self._setup_styles()
        self._add_header_footer(system_name)

        self._add_title_page(system_name, profile)
        self._add_summary_page(scan_result.stats, scan_result.total_folders, profile)

        files_by_part = self._group_files(scan_result)

        toc_entries = self._calculate_toc(files_by_part, profile)
        self._add_toc(toc_entries)

        self._write_parts(files_by_part, profile, scan_result)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save to a temporary file first. This prevents a failed safety check
        # from replacing an existing document with an unsafe one.
        temp_fd, temp_name = tempfile.mkstemp(
            prefix=f".{output_path.stem}-",
            suffix=output_path.suffix or ".docx",
            dir=output_path.parent,
        )
        os.close(temp_fd)
        temp_path = Path(temp_name)

        try:
            self.doc.save(str(temp_path))
            self._validate_docx_secrets(temp_path)
            os.replace(temp_path, output_path)
        finally:
            if temp_path.exists():
                temp_path.unlink()

        logger.info(f"Document saved: {output_path}")
